"""
简历-JD 智能匹配打分工具

用法:
  # JD 管理
  python jd_matcher.py --add-jd "AI全栈工程师" --jd-text "岗位职责：..."
  python jd_matcher.py --add-jd "AI全栈工程师" --jd-file "jd.txt"
  python jd_matcher.py --list-jd
  python jd_matcher.py --remove-jd "AI全栈工程师"

  # 简历打分
  python jd_matcher.py --score "d:/简历/张三_AI全栈工程师.pdf"
  python jd_matcher.py --score "简历.docx" --jd "AI全栈工程师"
  python jd_matcher.py --match "d:/简历/"              # 批量匹配整个文件夹
  python jd_matcher.py --match "d:/简历/" --output csv  # 批量导出CSV汇总

环境变量:
  DEEPSEEK_API_KEY    DeepSeek API Key (必需)
  DEEPSEEK_BASE_URL   DeepSeek API 地址 (默认 https://api.deepseek.com)
"""

import argparse
import csv
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

# ============================================================
# 环境变量加载
# ============================================================
env_path = Path(__file__).parent / ".env"
if env_path.exists():
    for line in env_path.read_text("utf-8").splitlines():
        line = line.strip()
        if line and not line.startswith("#") and "=" in line:
            key, val = line.split("=", 1)
            os.environ.setdefault(key.strip(), val.strip())

DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "")
DEEPSEEK_BASE_URL = os.environ.get("DEEPSEEK_BASE_URL", "https://api.deepseek.com")

# 路径
SCRIPT_DIR = Path(__file__).parent
JDS_FILE = SCRIPT_DIR / "jds.json"
REPORTS_DIR = SCRIPT_DIR / "match_reports"

# ============================================================
# JD 管理
# ============================================================

def load_jds() -> dict:
    """加载所有JD，返回 {title: {...}} 字典"""
    if not JDS_FILE.exists():
        return {}
    return json.loads(JDS_FILE.read_text("utf-8"))


def save_jds(jds: dict):
    JDS_FILE.write_text(json.dumps(jds, ensure_ascii=False, indent=2), "utf-8")


def add_jd(title: str, jd_text: str):
    """添加或更新一个JD"""
    jds = load_jds()
    now = datetime.now().strftime("%Y-%m-%dT%H:%M:%S")
    created_at = jds.get(title, {}).get("created_at", now)

    # 自动提取标签（从标题和JD中提取关键词）
    tags = extract_tags(title, jd_text)

    is_new = title not in jds
    jds[title] = {
        "title": title,
        "jd_text": jd_text.strip(),
        "tags": tags,
        "created_at": created_at,
        "updated_at": now,
    }
    save_jds(jds)
    action = "添加" if is_new else "更新"
    print(f"[OK] {action}JD: {title}")


def remove_jd(title: str):
    jds = load_jds()
    if title not in jds:
        print(f"[失败] 找不到岗位: {title}")
        return
    del jds[title]
    save_jds(jds)
    print(f"[OK] 已删除JD: {title}")


def list_jds():
    jds = load_jds()
    if not jds:
        print("[提示] 还没有添加任何JD，请先用 --add-jd 添加")
        return
    print(f"\n已录入 {len(jds)} 个岗位:\n")
    for i, (title, jd) in enumerate(jds.items(), 1):
        preview = jd["jd_text"][:80].replace("\n", " ")
        print(f"  {i}. {title}")
        print(f"     标签: {', '.join(jd.get('tags', []))}")
        print(f"     预览: {preview}...")
        print()


def extract_tags(title: str, jd_text: str) -> list:
    """从标题和JD中自动提取关键词标签"""
    combined = title + " " + jd_text[:500]
    # 常见技术/岗位关键词
    keyword_pool = [
        "Python", "Java", "Go", "Rust", "C++", "JavaScript", "TypeScript",
        "React", "Vue", "Angular", "Node.js", "Next.js",
        "AI", "机器学习", "深度学习", "NLP", "CV", "大模型", "LLM",
        "全栈", "前端", "后端", "DevOps", "MCP", "API",
        "招聘", "培训", "HRBP", "薪酬", "绩效", "组织发展",
        "MySQL", "PostgreSQL", "MongoDB", "Redis",
        "Docker", "Kubernetes", "AWS", "Azure", "GCP",
        "敏捷", "Scrum", "项目管理",
        "数据分析", "数据工程", "爬虫",
    ]
    tags = []
    for kw in keyword_pool:
        if kw.lower() in combined.lower():
            tags.append(kw)
    # 最多取8个
    return tags[:8]


# ============================================================
# 简历解析
# ============================================================

def parse_resume(file_path: str) -> tuple:
    """
    解析简历文件，返回 (文本内容, 候选人姓名推测)
    支持: PDF, DOCX, TXT, MD
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    suffix = path.suffix.lower()
    text = ""

    if suffix == ".pdf":
        text = _parse_pdf(path)
    elif suffix in (".docx", ".doc"):
        text = _parse_docx(path)
    elif suffix in (".txt", ".md", ".text"):
        text = path.read_text("utf-8", errors="replace")
    else:
        raise ValueError(f"不支持的文件格式: {suffix}（支持 PDF/DOCX/TXT/MD）")

    if not text.strip():
        raise ValueError(f"无法从文件中提取文字: {file_path}")

    # 从文件名和文本推测候选人姓名
    name = _guess_name(path.stem, text)
    return text.strip(), name


def _parse_pdf(path: Path) -> str:
    """用 pdfplumber 解析 PDF"""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("需要安装 pdfplumber: pip install pdfplumber")

    text_parts = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def _parse_docx(path: Path) -> str:
    """用 python-docx 解析 Word 文档"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("需要安装 python-docx: pip install python-docx")

    doc = Document(str(path))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # 同时提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
            if row_text.strip():
                text_parts.append(row_text)

    return "\n".join(text_parts)


def _guess_name(filename_stem: str, text: str) -> str:
    """从文件名或简历文本首行推测候选人姓名"""
    # 优先从文件名提取：常见格式为 "姓名_岗位" 或 "简历_姓名"
    stem = filename_stem
    # 去掉常见前缀
    for prefix in ["简历_", "resume_", "个人简历_", "Resume_", "简历-", "个人简历-"]:
        if stem.lower().startswith(prefix.lower()):
            stem = stem[len(prefix):]

    # 取第一个下划线或连字符前的部分作为姓名
    for sep in ["_", "－", "-", " "]:
        if sep in stem:
            candidate = stem.split(sep)[0].strip()
            if 2 <= len(candidate) <= 5 and not any(kw in candidate.lower() for kw in ["ai", "hr", "工程师", "经理", "开发"]):
                return candidate

    # 如果文件名推断不出，看文本首行
    first_line = text.strip().split("\n")[0][:20]
    if len(first_line) <= 5 and not any(kw in first_line.lower() for kw in ["简历", "resume", "个人"]):
        return first_line

    return "未知"


# ============================================================
# 文件名 → JD 匹配
# ============================================================

def match_jd_from_filename(filename: str, jds: dict) -> tuple:
    """
    根据简历文件名匹配最合适的JD。
    返回 (jd_title, score) 或 (None, 0) 如果没有匹配。
    """
    if not jds:
        return None, 0

    stem = Path(filename).stem
    # 清理常见前缀
    for prefix in ["简历_", "resume_", "个人简历_", "Resume_", "简历-", "个人简历-"]:
        if stem.lower().startswith(prefix.lower()):
            stem = stem[len(prefix):]

    # 策略1: 直接子串匹配（如 "张三_AI全栈工程师" 包含 "AI全栈工程师"）
    stem_lower = stem.lower()
    for jd_title in jds:
        if jd_title.lower() in stem_lower:
            return jd_title, 1.0

    # 策略2: 分词匹配
    chars_to_split = r"[_\-\s]+"
    file_words = set()
    for part in re.split(chars_to_split, stem):
        part = part.strip().lower()
        if len(part) >= 2:
            file_words.add(part)

    if not file_words:
        return None, 0

    best_jd = None
    best_score = 0

    for jd_title in jds:
        # JD标题分词
        jd_title_words = set()
        for part in re.split(chars_to_split, jd_title.lower()):
            part = part.strip()
            if len(part) >= 2:
                jd_title_words.add(part)

        # 标签也参与匹配（扩大命中范围）
        tag_words = set(t.lower() for t in jds[jd_title].get("tags", []))

        if not jd_title_words:
            continue

        # 基于标题词计算匹配度（交集/标题词数）
        title_overlap = file_words & jd_title_words
        score = len(title_overlap) / len(jd_title_words)

        # 标签命中加分：每命中一个标签 +0.05，最多 +0.3
        tag_overlap = file_words & tag_words
        tag_bonus = min(len(tag_overlap) * 0.05, 0.3)
        score += tag_bonus

        if score > best_score:
            best_score = score
            best_jd = jd_title

    return best_jd, best_score


# ============================================================
# AI 评分（DeepSeek）
# ============================================================

SCORING_PROMPT = """你是一位资深的招聘专家和HR顾问。请根据以下岗位JD和候选人简历，进行专业评估打分。

## 评分规则
从以下六个维度打分，总分100分：

1. **技术技能匹配** (0-30分)：技术栈、工具、编程语言与岗位要求的匹配程度
2. **工作经验匹配** (0-25分)：工作年限、行业背景、岗位经验的匹配程度
3. **项目经验匹配** (0-20分)：项目复杂度、成果、与岗位的相关性
4. **学历/资质** (0-15分)：学历层次、专业方向、证书资质
5. **综合素质** (0-10分)：沟通表达、团队协作、学习能力（从简历描述中推断）
6. **总分** (0-100分)

## 输出要求
请严格返回以下JSON格式，不要有任何其他文字：

{{
  "overall_score": 85,
  "level": "推荐",
  "scores": {{
    "tech_skill": 26,
    "work_exp": 20,
    "project_exp": 17,
    "education": 12,
    "overall_quality": 8,
    "total": 83
  }},
  "highlights": ["精通 Python 和 React，技术栈高度匹配", "有 3 年 AI 项目经验"],
  "gaps": ["缺少 Kubernetes 运维经验", "学历稍低于要求"],
  "interview_suggestions": ["重点考察大模型调优实战经验", "了解其对全栈架构的理解深度"],
  "summary": "候选人技术基础扎实，项目经验与岗位需求匹配度较高，建议进入面试环节。"
}}

level 必须是以下四选一：
- "强烈推荐"（总分 >= 90）
- "推荐"（总分 >= 75）
- "待定"（总分 >= 60）
- "不推荐"（总分 < 60）

highlights: 2-4个匹配亮点
gaps: 1-3个差距点，如果没有明显差距可以写["无明显差距"]
interview_suggestions: 2-4条面试考察建议
summary: 50字以内的综合评价

---

## 岗位JD
{jd_text}

## 候选人简历
{resume_text}

---

请开始评估，只返回JSON。"""


def call_deepseek(prompt: str, max_tokens: int = 2000, use_json: bool = True) -> str:
    """调用 DeepSeek API"""
    if not DEEPSEEK_API_KEY:
        raise SystemExit("请设置环境变量 DEEPSEEK_API_KEY")

    body = {
        "model": "deepseek-chat",
        "messages": [
            {"role": "system", "content": "你是一位资深的招聘专家。只返回JSON，不要有其他内容。"},
            {"role": "user", "content": prompt}
        ],
        "max_tokens": max_tokens,
        "temperature": 0.3,
    }
    if use_json:
        body["response_format"] = {"type": "json_object"}

    from urllib.request import Request, urlopen

    req = Request(
        f"{DEEPSEEK_BASE_URL}/v1/chat/completions",
        data=json.dumps(body).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json",
        }
    )
    try:
        resp = urlopen(req, timeout=180)
        result = json.loads(resp.read().decode("utf-8"))
        return result["choices"][0]["message"]["content"]
    except Exception as e:
        raise SystemExit(f"DeepSeek API 调用失败: {e}")


def score_resume(resume_text: str, jd_text: str) -> dict:
    """用 AI 对简历进行打分，返回评分结果 dict"""
    # 限制长度节省 token
    jd_trimmed = jd_text[:5000] if len(jd_text) > 5000 else jd_text
    resume_trimmed = resume_text[:12000] if len(resume_text) > 12000 else resume_text

    prompt = SCORING_PROMPT.format(
        jd_text=jd_trimmed,
        resume_text=resume_trimmed,
    )
    result = call_deepseek(prompt, max_tokens=2000)
    try:
        return json.loads(result)
    except json.JSONDecodeError:
        # 尝试从结果中提取 JSON
        match = re.search(r'\{.*\}', result, re.DOTALL)
        if match:
            return json.loads(match.group())
        raise SystemExit(f"AI 返回结果解析失败:\n{result}")


# ============================================================
# 报告输出
# ============================================================

def print_score_report(jd_title: str, candidate_name: str, result: dict):
    """在控制台打印简洁评分报告"""
    scores = result.get("scores", {})
    total = scores.get("total", result.get("overall_score", 0))
    level = result.get("level", "未知")

    # 等级图标
    level_icons = {"强烈推荐": "★", "推荐": "☆", "待定": "○", "不推荐": "×"}
    icon = level_icons.get(level, "?")

    print()
    print("=" * 60)
    print(f"  简历评分报告")
    print(f"  岗位: {jd_title}    候选人: {candidate_name}")
    print(f"  总分: {total}/100  {icon} {level}")
    print("-" * 60)
    print(f"  技术技能: {scores.get('tech_skill', 0):>2}/30  工作经验: {scores.get('work_exp', 0):>2}/25")
    print(f"  项目经验: {scores.get('project_exp', 0):>2}/20  学历资质: {scores.get('education', 0):>2}/15")
    print(f"  综合素质: {scores.get('overall_quality', 0):>2}/10")
    print("-" * 60)

    highlights = result.get("highlights", [])
    if highlights:
        print("  [匹配亮点]")
        for h in highlights:
            print(f"    + {h}")

    gaps = result.get("gaps", [])
    if gaps and gaps != ["无明显差距"]:
        print("  [差距分析]")
        for g in gaps:
            print(f"    - {g}")

    suggestions = result.get("interview_suggestions", [])
    if suggestions:
        print("  [面试建议]")
        for s in suggestions:
            print(f"    → {s}")

    summary = result.get("summary", "")
    if summary:
        print(f"\n  📝 {summary}")
    print("=" * 60)
    print()


def save_json_report(jd_title: str, candidate_name: str, result: dict, file_path: str):
    """保存详细 JSON 报告"""
    REPORTS_DIR.mkdir(exist_ok=True)
    date_str = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = re.sub(r'[\\/:*?"<>|]', '_', candidate_name)
    safe_jd = re.sub(r'[\\/:*?"<>|]', '_', jd_title)
    report_file = REPORTS_DIR / f"{date_str}_{safe_jd}_{safe_name}.json"

    report = {
        "generated_at": datetime.now().strftime("%Y-%m-%dT%H:%M:%S"),
        "jd_title": jd_title,
        "candidate_name": candidate_name,
        "resume_file": str(file_path),
        "result": result,
    }
    report_file.write_text(json.dumps(report, ensure_ascii=False, indent=2), "utf-8")
    print(f"  详细报告已保存: {report_file}")


def export_csv(results: list, output_path: str = None):
    """批量匹配结果导出为 CSV"""
    if not output_path:
        date_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = REPORTS_DIR / f"batch_summary_{date_str}.csv"
    else:
        output_path = Path(output_path)

    REPORTS_DIR.mkdir(exist_ok=True)

    with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow([
            "候选人", "岗位", "总分", "推荐等级",
            "技术技能/30", "工作经验/25", "项目经验/20", "学历/15", "综合素质/10",
            "亮点", "差距", "面试建议", "综合评价", "简历文件"
        ])
        for r in results:
            scores = r["result"].get("scores", {})
            writer.writerow([
                r["candidate_name"],
                r["jd_title"],
                scores.get("total", r["result"].get("overall_score", 0)),
                r["result"].get("level", ""),
                scores.get("tech_skill", ""),
                scores.get("work_exp", ""),
                scores.get("project_exp", ""),
                scores.get("education", ""),
                scores.get("overall_quality", ""),
                "；".join(r["result"].get("highlights", [])),
                "；".join(r["result"].get("gaps", [])),
                "；".join(r["result"].get("interview_suggestions", [])),
                r["result"].get("summary", ""),
                r["resume_file"],
            ])
    print(f"  CSV汇总已保存: {output_path}")


# ============================================================
# 单份简历处理
# ============================================================

def process_single_resume(resume_path: str, jd_title: str = None):
    """处理单份简历：解析 → 匹配JD → 打分 → 输出报告"""
    path = Path(resume_path)

    # 1. 解析简历
    print(f"\n[1/4] 解析简历: {path.name}")
    resume_text, candidate_name = parse_resume(str(path))
    char_count = len(resume_text)
    print(f"      候选人: {candidate_name} | 提取 {char_count} 字符")

    # 2. 匹配 JD
    jds = load_jds()
    if not jds:
        raise SystemExit("[失败] 请先用 --add-jd 添加岗位JD")

    if jd_title:
        if jd_title not in jds:
            raise SystemExit(f"[失败] 找不到岗位「{jd_title}」，可用岗位: {', '.join(jds.keys())}")
        match_score = 1.0
        print(f"[2/4] 指定岗位: {jd_title}")
    else:
        print(f"[2/4] 自动匹配岗位...")
        jd_title, match_score = match_jd_from_filename(path.name, jds)

    if not jd_title or match_score < 0.2:
        print(f"      匹配度 {match_score:.0%}，无法自动确定岗位")
        print(f"      可用岗位: {', '.join(jds.keys())}")
        if sys.stdin.isatty():
            jd_title = input("      请手动输入岗位名: ").strip()
        else:
            raise SystemExit(
                "[失败] 无法自动匹配岗位，请用 --jd 参数指定岗位名。\n"
                f"      可用岗位: {', '.join(jds.keys())}"
            )
        if jd_title not in jds:
            raise SystemExit(f"[失败] 找不到岗位「{jd_title}」")
    else:
        print(f"      匹配到: {jd_title}（匹配度 {match_score:.0%}）")

    # 3. AI 打分
    jd_text = jds[jd_title]["jd_text"]
    print(f"[3/4] AI 评分中...")
    result = score_resume(resume_text, jd_text)

    # 4. 输出报告
    print(f"[4/4] 生成报告...")
    print_score_report(jd_title, candidate_name, result)
    save_json_report(jd_title, candidate_name, result, str(path))

    return {
        "candidate_name": candidate_name,
        "jd_title": jd_title,
        "resume_file": str(path),
        "result": result,
    }


# ============================================================
# 批量处理
# ============================================================

def process_folder(folder_path: str, output_format: str = None):
    """批量处理文件夹内的所有简历"""
    folder = Path(folder_path)
    if not folder.is_dir():
        raise SystemExit(f"文件夹不存在: {folder_path}")

    supported = {".pdf", ".docx", ".doc", ".txt", ".md"}
    files = [f for f in folder.iterdir() if f.suffix.lower() in supported]

    if not files:
        print(f"[提示] 文件夹内没有支持的简历文件（PDF/DOCX/TXT/MD）: {folder_path}")
        return

    jds = load_jds()
    if not jds:
        raise SystemExit("[失败] 请先用 --add-jd 添加岗位JD")

    print(f"\n找到 {len(files)} 份简历，开始批量处理...\n")
    print("=" * 60)

    all_results = []
    for i, file in enumerate(files, 1):
        try:
            print(f"\n[{i}/{len(files)}] {file.name}")
            jd_title, match_score = match_jd_from_filename(file.name, jds)

            if not jd_title or match_score < 0.2:
                print(f"      跳过：无法匹配到合适岗位（匹配度 {match_score:.0%}）")
                continue

            print(f"      匹配: {jd_title}（{match_score:.0%}）")

            resume_text, candidate_name = parse_resume(str(file))
            print(f"      候选人: {candidate_name} | {len(resume_text)} 字符")

            jd_text = jds[jd_title]["jd_text"]
            print(f"      AI 评分中...")
            result = score_resume(resume_text, jd_text)

            print_score_report(jd_title, candidate_name, result)
            save_json_report(jd_title, candidate_name, result, str(file))

            all_results.append({
                "candidate_name": candidate_name,
                "jd_title": jd_title,
                "resume_file": str(file),
                "result": result,
            })

        except Exception as e:
            print(f"      [错误] {e}")
            continue

    # 汇总
    print(f"\n{'=' * 60}")
    print(f"批量处理完成: {len(all_results)}/{len(files)} 份成功")
    if all_results:
        # 按总分排序
        all_results.sort(
            key=lambda r: r["result"].get("scores", {}).get("total", r["result"].get("overall_score", 0)),
            reverse=True
        )
        print("\n  排名汇总:")
        for i, r in enumerate(all_results, 1):
            scores = r["result"].get("scores", {})
            total = scores.get("total", r["result"].get("overall_score", 0))
            level = r["result"].get("level", "")
            print(f"  {i}. {r['candidate_name']:6s} | {r['jd_title']:12s} | {total:>3}分 | {level}")

        if output_format == "csv":
            export_csv(all_results)


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="简历-JD 智能匹配打分工具",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python jd_matcher.py --add-jd "AI全栈工程师" --jd-text "岗位职责：..."
  python jd_matcher.py --add-jd "AI全栈工程师" --jd-file "jd.txt"
  python jd_matcher.py --list-jd
  python jd_matcher.py --remove-jd "AI全栈工程师"
  python jd_matcher.py --score "d:/简历/张三_AI全栈工程师.pdf"
  python jd_matcher.py --score "简历.docx" --jd "AI全栈工程师"
  python jd_matcher.py --match "d:/简历/"
  python jd_matcher.py --match "d:/简历/" --output csv
        """
    )

    # JD 管理
    parser.add_argument("--add-jd", metavar="岗位名", help="添加或更新岗位JD")
    parser.add_argument("--jd-text", metavar="JD内容", help="JD文本内容（配合 --add-jd）")
    parser.add_argument("--jd-file", metavar="文件路径", help="从文件读取JD内容（配合 --add-jd）")
    parser.add_argument("--list-jd", action="store_true", help="列出所有已录入的岗位JD")
    parser.add_argument("--remove-jd", metavar="岗位名", help="删除指定岗位JD")

    # 简历评分
    parser.add_argument("--score", metavar="简历文件", help="对单份简历进行打分")
    parser.add_argument("--jd", metavar="岗位名", help="指定JD岗位（配合 --score，不指定则自动匹配）")
    parser.add_argument("--match", metavar="文件夹", help="批量匹配文件夹内所有简历")
    parser.add_argument("--output", choices=["csv"], help="批量匹配时导出CSV汇总")

    args = parser.parse_args()

    # ---- JD 管理命令 ----
    if args.add_jd:
        # 获取 JD 文本
        if args.jd_text:
            jd_text = args.jd_text
        elif args.jd_file:
            jd_path = Path(args.jd_file)
            if not jd_path.exists():
                raise SystemExit(f"JD文件不存在: {args.jd_file}")
            jd_text = jd_path.read_text("utf-8", errors="replace")
        else:
            # 交互式输入
            print(f"请输入「{args.add_jd}」的JD内容（输入完成后按 Ctrl+Z 然后回车结束）:")
            lines = sys.stdin.read()
            jd_text = lines.strip()
            if not jd_text:
                raise SystemExit("[失败] JD内容不能为空")

        add_jd(args.add_jd, jd_text)
        return

    if args.list_jd:
        list_jds()
        return

    if args.remove_jd:
        remove_jd(args.remove_jd)
        return

    # ---- 简历评分命令 ----
    if args.score:
        process_single_resume(args.score, args.jd)
        return

    if args.match:
        process_folder(args.match, args.output)
        return

    # 无参数时打印帮助
    parser.print_help()


if __name__ == "__main__":
    main()
