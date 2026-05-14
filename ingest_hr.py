"""
HR知识库 · 内容入库工具（全功能版）

两种使用模式：
  1. 日常CLI模式：
     python ingest_hr.py --file "d:/资料/xxx.pdf"
     python ingest_hr.py --folder "d:/HR资料/"
     python ingest_hr.py --url "https://www.xiaohongshu.com/..."

  2. 后台自动处理模式（Base为主）：
     python ingest_hr.py --daemon           # 常驻轮询，自动处理Base待处理记录
     python ingest_hr.py --sync             # 单次处理，适合Windows计划任务

做法：在Base创建记录 → 填"原始链接" → 保存 → 后台自动抓取/AI分析/填充结果

选项:
  --review       强制交互式预览（默认：单文件/文本/链接时开启，批量时关闭）
  --yes          跳过所有确认（用于批量/定时任务）
  --title        手动指定标题（配合 --text 使用）
  --poll-interval Daemon轮询间隔秒数（默认60）

环境变量:
  DEEPSEEK_API_KEY         DeepSeek API Key (必需)
  DEEPSEEK_BASE_URL        DeepSeek API 地址 (默认 https://api.deepseek.com)

  小红书登录状态:
  PLAYWRIGHT_STATE_DIR      Playwright持久化状态目录（首次自动引导登录）

  语音识别 (视频转文字):
  ASR_ENGINE               引擎: lark(飞书妙记/默认) | whisper
  飞书妙记: 无需配置，自动上传音频到云空间→生成妙记→获取逐字稿
  Whisper:  需 WHISPER_API_KEY + WHISPER_BASE_URL
"""

import argparse, json, os, re, shutil, subprocess, sys, tempfile, time, signal
from datetime import datetime
from pathlib import Path
from urllib.request import Request, urlopen
from urllib.parse import urlparse

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

# 读取 .env 文件（如有）
env_path = Path(__file__).parent / ".env"
if env_path.exists():
    for line in env_path.read_text("utf-8").splitlines():
        line = line.strip()
        if line and not line.startswith("#") and "=" in line:
            key, val = line.split("=", 1)
            os.environ.setdefault(key.strip(), val.strip())

# ============================================================
# 配置
# ============================================================
BASE_TOKEN = "TpoSbBr6QaXUDFs4abYcEJQEnUd"
TABLE_ID = "tblac5PrdBNBj8Nn"

DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "")
DEEPSEEK_BASE_URL = os.environ.get("DEEPSEEK_BASE_URL", "https://api.deepseek.com")

WHISPER_API_KEY = os.environ.get("WHISPER_API_KEY", "")
WHISPER_BASE_URL = os.environ.get("WHISPER_BASE_URL", "https://api.openai.com/v1")

# 国内语音识别（优先使用 FunASR 本地模型，也可配置云 API）
BAIDU_ASR_APP_ID = os.environ.get("BAIDU_ASR_APP_ID", "")
BAIDU_ASR_API_KEY = os.environ.get("BAIDU_ASR_API_KEY", "")
BAIDU_ASR_SECRET_KEY = os.environ.get("BAIDU_ASR_SECRET_KEY", "")
ALIYUN_ASR_APP_KEY = os.environ.get("ALIYUN_ASR_APP_KEY", "")
ALIYUN_ASR_ACCESS_KEY_ID = os.environ.get("ALIYUN_ASR_ACCESS_KEY_ID", "")
ALIYUN_ASR_ACCESS_KEY_SECRET = os.environ.get("ALIYUN_ASR_ACCESS_KEY_SECRET", "")

ASR_ENGINE = os.environ.get("ASR_ENGINE", "lark")  # lark(飞书妙记) | whisper

# Playwright 持久化登录状态目录（小红书）
PLAYWRIGHT_STATE_DIR = Path(os.environ.get("PLAYWRIGHT_STATE_DIR",
    str(Path.home() / ".hr_ingest" / "playwright_state")))
PLAYWRIGHT_STATE_FILE = PLAYWRIGHT_STATE_DIR / "auth_state.json"

# Daemon 配置
DAEMON_POLL_INTERVAL = int(os.environ.get("DAEMON_POLL_INTERVAL", "60"))

TAG_POOL = [
    # 职场成长
    "沟通表达", "效率工具", "职业规划", "职场关系",
    # HR知识成长
    "招聘面试", "培训发展", "薪酬绩效", "员工关系", "组织发展", "劳动法务",
]

SOURCE_POOL = ["视频转写", "网页提取", "小红书", "飞书文档", "手动录入"]

USER_AGENT = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"

# ============================================================
# 文件解析
# ============================================================
def read_txt(path: str) -> str:
    for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()

def read_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)

def read_pdf(path: str) -> str:
    import fitz
    doc = fitz.open(path)
    parts = [page.get_text() for page in doc]
    doc.close()
    return "\n\n".join(parts)

def extract_text(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    parsers = {".txt": read_txt, ".md": read_txt, ".docx": read_docx, ".pdf": read_pdf}
    parser = parsers.get(ext)
    if not parser:
        raise ValueError(f"不支持的文件类型: {ext}")
    return parser(file_path)

# ============================================================
# 网页抓取
# ============================================================
def fetch_web_content(url: str) -> str:
    """从网页URL提取正文。先试 trafilatura，再试 readability，最后用 requests + lxml"""
    print("  🌐 抓取网页内容...")

    # 方案1: trafilatura（最准，特别是文章类）
    try:
        import trafilatura
        downloaded = trafilatura.fetch_url(url)
        if downloaded:
            text = trafilatura.extract(downloaded, include_comments=False, include_tables=True)
            if text and len(text) > 100:
                print(f"  ✅ trafilatura 提取成功 ({len(text)} 字符)")
                return text
    except Exception:
        pass

    # 方案2: readability-lxml + requests
    try:
        from readability import Document
        import requests
        headers = {"User-Agent": USER_AGENT}
        resp = requests.get(url, headers=headers, timeout=30)
        resp.encoding = resp.apparent_encoding
        doc = Document(resp.text)
        html = doc.summary()
        from lxml import html as lxml_html
        tree = lxml_html.fromstring(html)
        text = tree.text_content().strip()
        if text and len(text) > 100:
            print(f"  ✅ readability 提取成功 ({len(text)} 字符)")
            return text
    except Exception:
        pass

    # 方案3: 纯 requests + lxml 提取所有文本
    try:
        import requests
        from lxml import html as lxml_html
        headers = {"User-Agent": USER_AGENT}
        resp = requests.get(url, headers=headers, timeout=30)
        resp.encoding = resp.apparent_encoding
        tree = lxml_html.fromstring(resp.text)
        # 移除脚本和样式
        for el in tree.xpath("//script | //style | //nav | //footer | //header"):
            el.getparent().remove(el)
        text = tree.text_content().strip()
        text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
        if text and len(text) > 50:
            print(f"  ✅ 基础HTML提取成功 ({len(text)} 字符)")
            return text
    except Exception:
        pass

    raise ValueError("无法从该URL提取内容（可能需登录或页面为动态加载）")

# ============================================================
# 小红书内容提取（浏览器渲染 + 视频下载 + 持久化登录态）
# ============================================================
PLAYWRIGHT_STATE_SESSION = PLAYWRIGHT_STATE_DIR / "auth_state.json"


def ensure_playwright_login() -> bool:
    """
    弹出浏览器，引导用户登录小红书，保存登录态。

    无论是否有旧状态都会弹出，确保 session 有效。
    """
    PLAYWRIGHT_STATE_DIR.mkdir(parents=True, exist_ok=True)
    print("\n🔑 需要重新登录小红书")
    print("  即将打开浏览器，请扫码或用手机号登录")

    try:
        from playwright.sync_api import sync_playwright
        with sync_playwright() as p:
            browser = p.chromium.launch(
                channel="chrome", headless=False,
                args=_STEALTH_ARGS,
            )
            context = browser.new_context(
                user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36",
                viewport={"width": 1280, "height": 800},
                locale="zh-CN",
                timezone_id="Asia/Shanghai",
            )
            context.add_init_script(_STEALTH_INIT_SCRIPT)
            page = context.new_page()
            page.goto("https://www.xiaohongshu.com", wait_until="domcontentloaded")
            input("  登录完成后按 Enter 继续...")
            context.storage_state(path=str(PLAYWRIGHT_STATE_SESSION))
            browser.close()
        print(f"  ✅ 登录状态已保存到 {PLAYWRIGHT_STATE_SESSION}")
        return True
    except ImportError:
        print("  ⚠️ 需要 playwright: pip install playwright")
        return False
    except Exception as e:
        print(f"  ⚠️ 登录引导失败: {e}")
        return False


_STEALTH_ARGS = [
    "--disable-blink-features=AutomationControlled",
    "--disable-automation",
    "--no-sandbox",
    "--disable-features=ChromeWhatsNew",
]

_STEALTH_INIT_SCRIPT = """
// Hide automation traces
Object.defineProperties(navigator, {
    webdriver: { get: () => false },
    plugins: { get: () => [1,2,3,4,5] },
    languages: { get: () => ['zh-CN', 'zh', 'en'] },
});
window.chrome = { runtime: {} };
// Override permissions
const origQuery = window.navigator.permissions.query;
window.navigator.permissions.query = (p) => (
    p.name === 'notifications' ? Promise.resolve({state: 'denied'}) : origQuery(p)
);
// WebGL vendor override
const getExt = HTMLCanvasElement.prototype.getContext;
HTMLCanvasElement.prototype.getContext = function(...args) {
    const ctx = getExt.apply(this, args);
    if (ctx && args[0] === 'webgl') {
        const getExt2 = ctx.getExtension;
        ctx.getExtension = function(name) {
            if (name === 'WEBGL_debug_renderer_info') return null;
            return getExt2.call(this, name);
        };
    }
    return ctx;
};
"""

# Playwright 无头浏览器额外参数


def get_playwright_context(playwright):
    """
    获取 Playwright 浏览器上下文，优先级：
    1. storage_state（持久化登录态）
    2. COOKIES_FILE（向后兼容）
    3. 无认证
    """
    launch_kwargs = dict(
        channel="chrome", headless=True,
        args=_STEALTH_ARGS,
    )

    # 方案1: storage_state
    if PLAYWRIGHT_STATE_SESSION.exists():
        try:
            context = playwright.chromium.launch(**launch_kwargs).new_context(
                storage_state=str(PLAYWRIGHT_STATE_SESSION),
                user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36",
                viewport={"width": 1920, "height": 1080},
                locale="zh-CN",
                timezone_id="Asia/Shanghai",
            )
            context.add_init_script(_STEALTH_INIT_SCRIPT)
            return context
        except Exception as e:
            print(f"  ⚠️ storage_state 加载失败: {e}")

    # 方案2: cookie 文件向后兼容
    cookies_file = os.environ.get("COOKIES_FILE", "")
    if cookies_file and os.path.exists(cookies_file):
        try:
            with open(cookies_file, "r", encoding="utf-8") as f:
                content = f.read().strip()
            cookies = json.loads(content) if content.startswith("[") else []
            if cookies:
                context = playwright.chromium.launch(**launch_kwargs).new_context(
                    user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
                    viewport={"width": 1920, "height": 1080},
                    locale="zh-CN",
                )
                context.add_init_script(_STEALTH_INIT_SCRIPT)
                for c in cookies:
                    try:
                        context.add_cookies([{
                            "name": c["name"], "value": c["value"],
                            "domain": c.get("domain", ".xiaohongshu.com"),
                            "path": c.get("path", "/"),
                        }])
                    except Exception:
                        pass
                print(f"  🍪 使用 cookie 文件 ({len(cookies)} 个)")
                return context
        except Exception:
            pass

    # 方案3: 无认证
    context = playwright.chromium.launch(**launch_kwargs).new_context(
        user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        viewport={"width": 1920, "height": 1080},
        locale="zh-CN",
    )
    context.add_init_script(_STEALTH_INIT_SCRIPT)
    return context


def fetch_xiaohongshu_content(url: str) -> tuple[str, str]:
    """
    提取小红书笔记内容。

    流程：
    1. 尝试 curl_cffi + 已保存 cookies → 从 SSR 的 __INITIAL_STATE__ 提取
    2. 如果未登录 → 引导用户重新登录 → 重试
    """
    print("  📱 小红书内容提取...")

    note_id = ""
    m_id = re.search(r'/item/([a-zA-Z0-9]+)', url)
    if m_id:
        note_id = m_id.group(1)

    # ---------- 用 curl_cffi + cookies 提取 ----------
    text, video_url = _fetch_xhs_via_curl(url, note_id)
    if text:
        return text, video_url

    # ---------- 未登录，引导重新登录 ----------
    print("  🔑 登录态已失效，需要重新登录")
    if not sys.stdin.isatty():
        print("  ⚠️ 请先运行: python ingest_hr.py --sync 在交互终端完成登录")
        return ("", "")

    if ensure_playwright_login():
        text, video_url = _fetch_xhs_via_curl(url, note_id)
        if text:
            return text, video_url

    return ("", "")


def _extract_initial_state(html: str) -> dict | None:
    """从 小红书 HTML 中提取 __INITIAL_STATE__ JSON"""
    p = "__INITIAL_STATE__="
    start = html.find(p)
    if start < 0:
        return None
    start += len(p)

    bc = 0
    in_str = False
    i = start
    while i < len(html):
        ch = html[i]
        if in_str:
            if ch == "\\":
                i += 2
                continue
            elif ch == '"':
                in_str = False
        else:
            if ch == '"':
                in_str = True
            elif ch == "{":
                bc += 1
            elif ch == "}":
                bc -= 1
                if bc == 0:
                    raw = html[start:i+1]
                    raw = re.sub(r":undefined", ":null", raw)
                    raw = re.sub(r",\s*}", "}", raw)
                    return json.loads(raw)
        i += 1
    return None


def _fetch_xhs_via_curl(url: str, note_id: str) -> tuple[str, str]:
    """用 curl_cffi 请求小红书页面，从 SSR 数据中提取笔记内容"""
    try:
        from curl_cffi import requests
    except ImportError:
        return ("", "")

    # 加载 cookies
    state_file = PLAYWRIGHT_STATE_SESSION
    if not state_file.exists():
        return ("", "")

    try:
        state = json.load(open(state_file, encoding="utf-8"))
    except Exception:
        return ("", "")

    cookies = {c["name"]: c["value"] for c in state.get("cookies", [])
               if "xiaohongshu.com" in c.get("domain", "")}
    if not cookies:
        return ("", "")

    cookie_header = "; ".join(f"{k}={v}" for k, v in cookies.items())
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36",
        "Cookie": cookie_header,
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
        "Referer": "https://www.xiaohongshu.com/",
    }

    try:
        r = requests.get(url, headers=headers, impersonate="chrome", timeout=15)
        if r.status_code != 200:
            print(f"  ⚠️ 请求失败: HTTP {r.status_code}")
            return ("", "")
    except Exception as e:
        print(f"  ⚠️ 请求异常: {e}")
        return ("", "")

    data = _extract_initial_state(r.text)
    if not data:
        print("  ⚠️ 未找到 __INITIAL_STATE__")
        return ("", "")

    logged_in = data.get("user", {}).get("loggedIn", False)

    # 从 noteDetailMap 提取笔记内容
    ndm = data.get("note", {}).get("noteDetailMap", {})
    for key in ndm:
        entry = ndm[key] if isinstance(ndm, dict) else {}
        note = entry.get("note", {}) if isinstance(entry, dict) else {}
        if note:
            desc = note.get("desc", "") or note.get("title", "") or note.get("content", "")
            if desc and isinstance(desc, str) and len(desc) > 10:
                # 提取视频信息
                video_url = ""
                video = note.get("video", {})
                if video:
                    for vk in ("stream", "media", "video_info", "original_info"):
                        info = video.get(vk, {}) or {}
                        if isinstance(info, dict):
                            streams = info.get("stream_list", []) or info.get("video_list", [])
                            for s in streams:
                                vu = s.get("url", s.get("video_url", ""))
                                if vu and ".mp4" in vu:
                                    video_url = vu
                                    break
                        if video_url:
                            break
                    if not video_url:
                        # 尝试直接取 media.url
                        media = video.get("media", {})
                        if isinstance(media, dict):
                            vu = media.get("url", "") or media.get("stream", {}).get("url", "")
                            if vu and ".mp4" in vu:
                                video_url = vu

                print(f"  ✅ SSR 提取到文字 ({len(desc)} 字符)")
                if video_url:
                    print(f"  🎬 发现视频: {video_url[:60]}...")
                return (desc.strip(), video_url)

    if not logged_in:
        print("  ⚠️ 登录态无效，需要重新登录")
        # 清理失效的 auth_state
        try:
            state_file.unlink()
        except Exception:
            pass

    return ("", "")

def ingest_xiaohongshu(full_url: str, review: bool, skip_confirm: bool) -> bool:
    """小红书全流程：提取文字 → 下载视频转写 → AI分析 → 入库"""
    print(f"\n📕 小红书笔记: {full_url}")

    # 1. 提取页面文字和视频URL
    text, video_url = fetch_xiaohongshu_content(full_url)

    # 2. 如果有视频，下载并转写
    transcription = ""
    if video_url:
        print("  🎬 发现视频，开始下载+转写...")
        try:
            transcription = transcribe_video(video_url)
        except Exception as e:
            print(f"  ⚠️ 视频处理失败: {e}")

    # 3. 合并文字和转写
    combined = text
    if transcription:
        separator = "\n\n=== 视频转写内容 ===\n\n"
        combined = text + separator + transcription if text else transcription

    if not combined.strip():
        print("  ❌ 未能获取任何内容")
        return False

    print(f"  📊 总 {len(combined)} 字符 (文字: {len(text)}, 视频转写: {len(transcription)})")

    # 4. 正常流程：AI分析 → 预览 → 查重 → 入库
    return process_and_write(
        title_hint="", text=combined, source="小红书", url_link=full_url,
        review=review, skip_confirm=skip_confirm
    )

# ============================================================
# 视频 → 音频 → 文字（Whisper API）
# ============================================================
def download_audio_from_video(url: str) -> str:
    """下载视频并提取音频，返回音频文件路径"""
    print("  📥 下载视频并提取音频...")
    tmpdir = tempfile.mkdtemp(prefix="hr_ingest_", dir=os.getcwd())
    output_template = os.path.join(tmpdir, "%(title)s.%(ext)s")

    try:
        import yt_dlp
        ydl_opts = {
            "format": "bestaudio/best",
            "postprocessors": [{
                "key": "FFmpegExtractAudio",
                "preferredcodec": "mp3",
                "preferredquality": "192",
            }],
            "outtmpl": output_template,
            "quiet": True,
            "no_warnings": True,
        }
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=True)
            title = info.get("title", "audio")
            # 查找生成的 mp3 文件
            for f in os.listdir(tmpdir):
                if f.endswith(".mp3"):
                    audio_path = os.path.join(tmpdir, f)
                    print(f"  ✅ 音频提取完成: {title}")
                    return audio_path
            raise FileNotFoundError("未找到提取的音频文件")
    except Exception as e:
        shutil.rmtree(tmpdir, ignore_errors=True)
        raise RuntimeError(f"视频下载失败: {e}")

def transcribe_via_lark_minutes(audio_path: str) -> str:
    """通过飞书妙记上传音频并获取逐字稿"""
    print("  🎙️ 飞书妙记转写...")

    # lark-cli 要求相对路径，转一下
    audio_rel = os.path.relpath(audio_path, os.getcwd())
    if not audio_rel.startswith("."):
        audio_rel = ".\\" + audio_rel if sys.platform == "win32" else "./" + audio_rel

    # 1. 上传音频到云空间
    print("  📤 上传音频到云空间...")
    resp = _run_lark_cli(["drive", "+upload", "--file", audio_rel], timeout=60)
    if not resp.get("ok"):
        raise RuntimeError(f"上传失败: {resp.get('error', {}).get('message', '未知')}")
    file_token = resp.get("data", {}).get("file_token", "")
    if not file_token:
        raise RuntimeError("未获取到 file_token")
    print(f"  ✅ 上传完成: {file_token[:12]}...")

    # 2. 生成妙记
    print("  🎬 生成妙记...")
    resp = _run_lark_cli(["minutes", "+upload", "--as", "user", "--file-token", file_token], timeout=30)
    if not resp.get("ok"):
        raise RuntimeError(f"生成妙记失败: {resp.get('error', {}).get('message', '未知')}")
    minute_url = resp.get("data", {}).get("minute_url", "")
    if not minute_url:
        raise RuntimeError("未获取到 minute_url")

    # 3. 提取 minute_token
    import re
    m = re.search(r'/minutes/([a-zA-Z0-9]+)', minute_url)
    if not m:
        raise RuntimeError(f"无法从URL提取 minute_token: {minute_url}")
    minute_token = m.group(1)
    print(f"  ✅ 妙记已生成: {minute_token[:12]}...")

    # 4. 等待妙记异步处理完成后获取逐字稿
    tmpdir = tempfile.mkdtemp(prefix="hr_ingest_minutes_")
    print("  ⏳ 等待转写完成...")
    for attempt in range(12):  # 最多等 2 分钟
        resp = _run_lark_cli([
            "vc", "+notes",
            "--as", "user",
            "--minute-tokens", minute_token,
            "--output-dir", tmpdir,
        ], timeout=30)
        if resp.get("ok"):
            break
        error_msg = resp.get("error", {}).get("message", "")
        if "processing" not in error_msg.lower() and "not ready" not in error_msg.lower():
            raise RuntimeError(f"获取逐字稿失败: {error_msg}")
        time.sleep(10)
    else:
        raise RuntimeError("妙记转写超时（2分钟）")

    # 5. 读取逐字稿文件
    transcript_dir = os.path.join(tmpdir, "minutes", minute_token)
    transcript_file = os.path.join(transcript_dir, "transcript.txt")
    if not os.path.exists(transcript_file):
        # 尝试旧布局
        for root, _, files in os.walk(tmpdir):
            for f in files:
                if f == "transcript.txt":
                    transcript_file = os.path.join(root, f)
                    break
    if not os.path.exists(transcript_file):
        raise RuntimeError("未找到逐字稿文件")

    with open(transcript_file, "r", encoding="utf-8") as f:
        text = f.read().strip()

    # 清理临时目录
    shutil.rmtree(tmpdir, ignore_errors=True)

    if text:
        print(f"  ✅ 转写完成 ({len(text)} 字符)")
        return text
    raise RuntimeError("逐字稿为空")


def transcribe_audio(audio_path: str) -> str:
    """将音频转为文字，默认使用飞书妙记"""
    engine = ASR_ENGINE

    # 飞书妙记（默认，推荐）
    if engine == "funasr" or engine == "lark":
        return transcribe_via_lark_minutes(audio_path)

    # Whisper API（远程备用）
    if engine == "whisper":
        api_key = WHISPER_API_KEY or DEEPSEEK_API_KEY
        if not api_key:
            raise SystemExit("Whisper API 需要设置 WHISPER_API_KEY")
        print("  🎙️ 语音转文字 (Whisper API)...")
        try:
            import requests
            with open(audio_path, "rb") as f:
                files = {"file": (os.path.basename(audio_path), f, "audio/mpeg")}
                data = {"model": "whisper-1", "language": "zh", "response_format": "json"}
                headers = {"Authorization": f"Bearer {api_key}"}
                resp = requests.post(f"{WHISPER_BASE_URL}/audio/transcriptions",
                                     headers=headers, files=files, data=data, timeout=300)
                resp.raise_for_status()
                result = resp.json()
                text = result.get("text", "").strip()
                if text:
                    print(f"  ✅ Whisper 转写完成 ({len(text)} 字符)")
                    return text
                raise ValueError("Whisper 返回空结果")
        except ImportError:
            raise SystemExit("需要 requests: pip install requests")
        except requests.exceptions.RequestException as e:
            raise RuntimeError(f"Whisper API 调用失败: {e}")

    raise SystemExit(f"未知的 ASR 引擎: {engine}，可选: lark, whisper")

def transcribe_video(url: str) -> str:
    """视频URL全流程：下载→提音频→转文字"""
    audio_path = None
    try:
        audio_path = download_audio_from_video(url)
        text = transcribe_audio(audio_path)
        return text
    finally:
        if audio_path:
            tmpdir = os.path.dirname(audio_path)
            shutil.rmtree(tmpdir, ignore_errors=True)

# ============================================================
# 飞书文档导入
# ============================================================
def fetch_feishu_doc(url: str) -> str:
    """从飞书文档/知识库链接提取内容"""
    print("  📄 获取飞书文档内容...")

    # 提取 token
    parsed = urlparse(url)
    path_parts = [p for p in parsed.path.split("/") if p]
    doc_token = None
    for p in path_parts:
        if len(p) > 10 and not p.startswith("~"):
            doc_token = p
            break

    if not doc_token:
        raise ValueError(f"无法从URL提取文档token: {url}")

    # 判断是否是 wiki 链接，先解析
    is_wiki = "wiki" in parsed.path
    if is_wiki:
        print(f"  🔗 Wiki 节点，先解析...")
        try:
            resp = _run_lark_cli(["wiki", "spaces", "get_node", "--node-token", doc_token], timeout=15)
            node = resp.get("data", {}).get("node", {})
            obj_token = node.get("obj_token", "") or node.get("obj_token", doc_token)
            obj_type = node.get("obj_type", "")
            if obj_type == "docx":
                doc_token = obj_token
            elif obj_type == "bitable":
                raise ValueError("Wiki 节点是多维表格，请直接使用 Base 链接")
            print(f"  ✅ 解析到文档 token: {doc_token}")
        except Exception as e:
            print(f"  ⚠️ Wiki 解析失败 ({e})，尝试直接用 token...")

    # 调用 lark-cli docs +fetch
    try:
        data = _run_lark_cli(["docs", "+fetch", "--doc", doc_token])

        # 优先取 markdown 字段（v2 API）或从 block 树提取
        markdown = data.get("data", {}).get("markdown", "")
        if not markdown or len(markdown.strip()) < 20:
            # 回退：从 block 树中提取所有文本
            texts = []
            def extract_blocks(blocks):
                if isinstance(blocks, list):
                    for block in blocks:
                        extract_blocks(block)
                elif isinstance(blocks, dict):
                    for key, val in blocks.items():
                        if key == "text" and isinstance(val, dict):
                            for item in val.get("elements", []):
                                if isinstance(item, dict):
                                    texts.append(item.get("text_run", {}).get("content", ""))
                        elif key == "title" and isinstance(val, dict):
                            for item in val.get("elements", []):
                                if isinstance(item, dict):
                                    texts.append(item.get("text_run", {}).get("content", ""))
                        elif key == "content":
                            extract_blocks(val)
                        elif isinstance(val, (dict, list)):
                            extract_blocks(val)
            extract_blocks(data)
            full_text = "\n".join(t for t in texts if t.strip())
            if full_text:
                print(f"  ✅ 飞书文档提取成功 ({len(full_text)} 字符)")
                return full_text
        else:
            print(f"  ✅ 飞书文档提取成功 ({len(markdown)} 字符)")
            return markdown
    except Exception as e:
        print(f"  ⚠️ docs +fetch 失败: {e}")

    raise ValueError("无法提取飞书文档内容")

# ============================================================
# DeepSeek API
# ============================================================
def call_deepseek(prompt: str, max_tokens: int = 1200) -> str:
    if not DEEPSEEK_API_KEY:
        raise SystemExit("请设置环境变量 DEEPSEEK_API_KEY")

    data = json.dumps({
        "model": "deepseek-chat",
        "messages": [
            {"role": "system", "content": "你是HR知识管理助手。只返回JSON，不要有其他内容。"},
            {"role": "user", "content": prompt}
        ],
        "max_tokens": max_tokens,
        "temperature": 0.3,
        "response_format": {"type": "json_object"}
    }).encode("utf-8")

    req = Request(f"{DEEPSEEK_BASE_URL}/v1/chat/completions", data=data, headers={
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    })
    try:
        resp = urlopen(req, timeout=120)
        result = json.loads(resp.read().decode("utf-8"))
        return result["choices"][0]["message"]["content"]
    except Exception as e:
        raise SystemExit(f"DeepSeek API 调用失败: {e}")

def analyze_hr_content(content: str) -> dict:
    """调用DeepSeek做内容整理：标题、统一格式化内容、自动分类打标签"""
    MAX_LEN = 8000
    if len(content) > MAX_LEN:
        content = content[:MAX_LEN] + "\n\n[... 内容过长，已截断]"

    prompt = f"""你是一个HR知识管理助手。请分析以下内容，返回JSON（不要有任何其他文字）：
{{
  "title": "内容标题（15字以内，概括核心主题）",
  "organized_content": "整理后的内容（格式统一、简洁精准，保留原文关键信息、要点、步骤、模板等，去除冗余和无关内容。用Markdown组织：标题用##、要点用-、步骤用1.）",
  "tags": ["标签1", "标签2"],
  "source_type": "来源类型"
}}

标签必须从以下列表选择1-3个。请精准判断内容属于哪个维度：
  职场成长（个人能力/职业发展）：沟通表达、效率工具、职业规划、职场关系
  HR知识成长（专业能力/业务知识）：招聘面试、培训发展、薪酬绩效、员工关系、组织发展、劳动法务
完整标签列表：{', '.join(TAG_POOL)}
来源类型必须从以下列表选择1个：{', '.join(SOURCE_POOL)}

=== 待分析内容 ===
{content}"""

    try:
        result = call_deepseek(prompt)
        parsed = json.loads(result)
        valid_tags = [t for t in parsed.get("tags", []) if t in TAG_POOL]
        valid_source = parsed.get("source_type", "") if parsed.get("source_type", "") in SOURCE_POOL else "手动录入"
        return {
            "title": parsed.get("title", "未命名"),
            "organized_content": parsed.get("organized_content", content[:500]),
            "tags": valid_tags[:3] or ["职业规划"],
            "source_type": valid_source,
        }
    except json.JSONDecodeError:
        return {
            "title": "未命名",
            "organized_content": content[:500],
            "tags": ["职业规划"],
            "source_type": "手动录入",
        }

# ============================================================
# 查重去重
# ============================================================
def check_duplicate(title: str, full_text: str = "", url_link: str = "") -> bool | None:
    """在Base中搜索相似内容。返回 True=找到重复，False=无重复，None=检查失败"""
    try:
        # 如果有链接，优先按原始链接查（最精确）
        if url_link:
            search_json = json.dumps({
                "keyword": url_link[:60],
                "search_fields": ["原始链接"],
                "limit": 3
            })
            data = _run_lark_cli(["base", "+record-search",
                "--base-token", BASE_TOKEN,
                "--table-id", TABLE_ID,
                "--json", search_json], timeout=15)
            items = data.get("data", {}).get("items", data.get("data", {}).get("records", []))
            if items:
                print(f"\n  ⚠️ 该链接已入库过，重复内容")
                return True

        # 按标题关键词搜索
        search_json = json.dumps({
            "keyword": title[:30],
            "search_fields": ["标题"],
            "limit": 5
        })
        data = _run_lark_cli(["base", "+record-search",
            "--base-token", BASE_TOKEN,
            "--table-id", TABLE_ID,
            "--json", search_json], timeout=15)

        records = []
        items = data.get("data", {}).get("items", data.get("data", {}).get("records", []))
        # 部分版本返回格式不同，尝试多种路径
        if not items and isinstance(data.get("data"), list):
            items = data["data"]
        if not items and isinstance(data, dict):
            for key in ("records", "items", "data"):
                val = data.get(key, data.get("data", {}).get(key, [])) if isinstance(data.get("data"), dict) else []
                if val:
                    items = val
                    break

        for item in items:
            if isinstance(item, dict):
                fields = item.get("fields", {})
                if not fields:
                    fields = {k: v for k, v in item.items() if k != "id"}
                existing_title = fields.get("标题", "")
                if existing_title and len(existing_title) > 2:
                    records.append(existing_title)

        if records:
            print(f"\n  ⚠️ 发现相似内容 ({len(records)} 条):")
            for r in records:
                print(f"     • {r[:50]}")
            return True
        return False
    except Exception as e:
        print(f"  ⚠️ 查重检查失败 ({e})，继续入库...")
        return None

# ============================================================
# 交互式预览
# ============================================================
def interactive_review(result: dict) -> dict | None:
    """用户预览和编辑AI分析结果。返回修改后的dict，或None表示跳过"""

    tags_menu = "\n".join(f"  [{i+1}] {t}" for i, t in enumerate(TAG_POOL))
    sources_menu = "\n".join(f"  [{i+1}] {s}" for i, s in enumerate(SOURCE_POOL))

    print("\n" + "═" * 55)
    print("  📋 AI 分析结果 — 请确认或修改")
    print("═" * 55)

    while True:
        print(f"""
  [1] 标题: {result['title']}
  [2] 来源: {result.get('source_type', '')}
  [3] 标签: {', '.join(result['tags'])}
  [4] 整理内容: {result['organized_content'][:200]}{'...' if len(result['organized_content']) > 200 else ''}""")

        print("  ────────────────────────────────────────")
        print("  [y] 确认入库  [s] 跳过  [数字] 修改对应字段")
        choice = input("  > ").strip().lower()

        if choice == "y":
            return result
        elif choice in ("s", "q", "n"):
            return None
        elif choice == "1":
            v = input(f"  标题 (回车不变: {result['title']}): ").strip()
            if v:
                result["title"] = v
        elif choice == "2":
            print(f"\n  可选来源:\n{sources_menu}")
            idx = input("  选择序号: ").strip()
            if idx.isdigit() and 1 <= int(idx) <= len(SOURCE_POOL):
                result["source_type"] = SOURCE_POOL[int(idx) - 1]
                print(f"  → 来源已更新: {result['source_type']}")
        elif choice == "3":
            print(f"\n  可选标签 (多选用逗号分隔):\n{tags_menu}")
            idxs = input("  选择序号: ").strip()
            selected = []
            for i in idxs.split(","):
                i = i.strip()
                if i.isdigit() and 1 <= int(i) <= len(TAG_POOL):
                    selected.append(TAG_POOL[int(i) - 1])
            if selected:
                result["tags"] = selected[:3]
                print(f"  → 标签已更新: {', '.join(result['tags'])}")
        elif choice == "4":
            print(f"  整理内容 (太长，仅显示前200字): {result['organized_content'][:200]}")
            v = input("  输入新内容替换 (回车不变):\n  > ").strip()
            if v:
                result["organized_content"] = v

# ============================================================
# lark-cli 调用封装
# ============================================================
# 直接调 Node.js run.js 绕过 cmd.exe 的参数截断和编码问题
_NPM_DIR = os.path.join(os.environ.get("APPDATA", ""), "npm")
_LARK_RUN_JS = os.path.join(_NPM_DIR, "node_modules", "@larksuite", "cli", "scripts", "run.js")
_NODE_EXE = shutil.which("node") or "node"


def _run_lark_cli(args: list[str], timeout: int = 30, max_retries: int = 2) -> dict:
    """直接调 Node.js 运行 lark-cli，避免 cmd.exe 编码问题"""
    if not os.path.exists(_LARK_RUN_JS):
        return {"ok": False, "error": {"message": f"lark-cli run.js 未找到: {_LARK_RUN_JS}"}}
    full_cmd = [_NODE_EXE, _LARK_RUN_JS] + args
    last_err = None
    # Windows 优先 GBK 再 UTF-8（中文系统输出是 CP936 编码）
    encoding_priority = ('gbk', 'cp936', 'utf-8') if sys.platform == 'win32' else ('utf-8', 'gbk', 'cp936')
    for attempt in range(max_retries + 1):
        try:
            result = subprocess.run(full_cmd, capture_output=True, timeout=timeout)
            # 多编码解析输出
            for enc in encoding_priority:
                try:
                    output = result.stdout.decode(enc)
                    if output.strip():
                        return json.loads(output)
                except (UnicodeDecodeError, json.JSONDecodeError):
                    continue
            # 尝试 stderr
            for enc in encoding_priority:
                try:
                    err_output = result.stderr.decode(enc)
                    if err_output.strip():
                        return json.loads(err_output)
                except (UnicodeDecodeError, json.JSONDecodeError):
                    continue
            last_err = f"Exit code {result.returncode}, no parseable output"
        except subprocess.TimeoutExpired:
            last_err = f"Timeout ({timeout}s)"
        except Exception as e:
            last_err = str(e)
        if attempt < max_retries:
            time.sleep(1)
    return {"ok": False, "error": {"message": last_err}}

# ============================================================
# 写入 HR知识库
# ============================================================
def write_to_hr_base(title: str, source: str, url_link: str, organized_content: str,
                     tags: list[str], raw_text: str = "") -> bool:
    record = {
        "标题": title,
        "来源": source,
        "整理内容": organized_content[:60000],
        "状态": "已完成",
    }
    if url_link:
        record["原始链接"] = url_link
    if tags:
        record["标签"] = tags
    if raw_text:
        record["原始内容"] = raw_text[:60000]

    body = json.dumps(record, ensure_ascii=False)

    # Windows 命令行长度限制 ~8000 字符，超出则先截断原始内容再截断整理内容
    MAX_CMD_LEN = 7500
    for field in ("原始内容", "整理内容"):
        if len(body) <= MAX_CMD_LEN:
            break
        content = record.get(field, "")
        if content:
            overhead = len(body) - len(content)
            available = MAX_CMD_LEN - overhead
            if available > 100:
                record[field] = content[:available]
                body = json.dumps(record, ensure_ascii=False)

    try:
        resp = _run_lark_cli([
            "base", "+record-upsert",
            "--base-token", BASE_TOKEN,
            "--table-id", TABLE_ID,
            "--json", body
        ])
        if resp.get("ok"):
            rec = resp.get("data", {}).get("record", {})
            rid = rec.get("id") or (rec.get("record_id_list", [None]) or [None])[0] or "?"
            print(f"  ✅ 已入库: {title} (record: {rid})")
            return True
        else:
            print(f"  ❌ 入库失败: {resp.get('error', {}).get('message', '未知')}")
            return False
    except Exception as e:
        print(f"  ❌ 调用 lark-cli 失败: {e}")
        return False


# ============================================================
# Base API 抽象层（Daemon/Sync 模式用）
# ============================================================
FIELD_状态 = "状态"
VALUE_待处理 = "待处理"
VALUE_处理中 = "处理中"
VALUE_已完成 = "已完成"
VALUE_失败 = "失败"


def _parse_record_list(resp: dict) -> list[dict]:
    """将 lark-cli +record-list --format json 的列式格式解析为记录列表"""
    data = resp.get("data", {})
    if isinstance(data, list):
        # 已经是行式格式
        return [{"id": r.get("id", ""), "fields": r.get("fields", r)} for r in data if isinstance(r, dict)]
    if not isinstance(data, dict):
        return []

    rows = data.get("data", [])
    field_names = data.get("fields", [])
    record_ids = data.get("record_id_list", [])

    records = []
    for i, row in enumerate(rows):
        if not isinstance(row, (list, tuple)):
            continue
        fields = {}
        for j, val in enumerate(row):
            if val is not None and j < len(field_names):
                name = field_names[j]
                # 处理 select 类型：单选用值数组，多选用值数组
                if isinstance(val, list) and len(val) > 0:
                    fields[name] = val[0] if len(val) == 1 else val
                else:
                    fields[name] = val
        rid = record_ids[i] if i < len(record_ids) else ""
        rec = {"id": rid}
        if fields:
            rec["fields"] = fields
        records.append(rec)
    return records


def base_get_pending_records(max_records: int = 20) -> list[dict]:
    """查询 Base 中所有状态=待处理的记录"""
    try:
        resp = _run_lark_cli([
            "base", "+record-list",
            "--base-token", BASE_TOKEN,
            "--table-id", TABLE_ID,
            "--limit", str(max_records),
            "--format", "json",
        ], timeout=20)
        records = _parse_record_list(resp)

        # 在 Python 侧过滤状态=待处理
        pending = []
        for r in records:
            fields = r.get("fields", {})
            status = fields.get(FIELD_状态, "")
            if status == VALUE_待处理 or not status:
                pending.append(r)
        return pending
    except Exception as e:
        print(f"  ⚠️ 查询待处理记录失败: {e}")
        return []


def base_update_record(record_id: str, fields: dict) -> bool:
    """更新指定记录的字段值（使用 +record-upsert --record-id）"""
    body = json.dumps(fields, ensure_ascii=False)
    try:
        resp = _run_lark_cli([
            "base", "+record-upsert",
            "--base-token", BASE_TOKEN,
            "--table-id", TABLE_ID,
            "--record-id", record_id,
            "--json", body,
        ], timeout=20)
        if resp.get("ok"):
            return True
        else:
            print(f"    ⚠️ 更新记录失败: {resp.get('error', {}).get('message', '未知')}")
            return False
    except Exception as e:
        print(f"    ⚠️ 更新记录异常: {e}")
        return False


def base_update_status(record_id: str, status: str) -> bool:
    """快捷更新状态"""
    fields = {FIELD_状态: status}
    return base_update_record(record_id, fields)


# ============================================================
# 记录处理路由（Daemon/Sync 模式用）
# ============================================================
_URL_EXTRACT_RE = re.compile(r'https?://[^\s）\)》\]]+')


def _clean_url(raw: str) -> str:
    """从可能的 markdown 链接或混合文本中提取纯 URL"""
    if not raw:
        return ""
    s = raw.strip()
    # 先尝试 markdown 格式 [text](url)，对 URL 用贪婪匹配拿到最后一个 )
    m = re.match(r'\[(.+?)\]\((.+)\)', s)
    if m:
        url_part = m.group(2).rstrip(")")
        # 如果 URL 部分本身是有效 URL（整段匹配，没有多余文字），直接使用
        if re.match(r'https?://[^\s）\)》\]]+$', url_part):
            return url_part
        # 否则在 text 部分找 https:// URL
        text_part = m.group(1)
        m2 = _URL_EXTRACT_RE.search(text_part)
        if m2:
            return m2.group(0).rstrip(")")
        # 再回退到 url_part 中找
        m2 = _URL_EXTRACT_RE.search(url_part)
        if m2:
            return m2.group(0).rstrip(")")
    # 直接从混合文本中提取 https:// 开头的 URL
    m = _URL_EXTRACT_RE.search(s)
    if m:
        return m.group(0).rstrip(")")
    return s


def process_base_record(record: dict) -> bool:
    """处理一条 Base 中的待处理记录，状态流转：待处理→处理中→已完成/失败"""
    record_id = record.get("id", "")
    fields = record.get("fields", {})
    raw_url = fields.get("原始链接", "") or fields.get("原始链接", {}).get("link", "")
    if raw_url and isinstance(raw_url, dict):
        raw_url = raw_url.get("link", "") or raw_url.get("text", "")
    url = _clean_url(raw_url)
    if not url:
        base_update_status(record_id, VALUE_失败)
        return False

    print(f"\n📝 处理记录 {record_id[:8]}...: {url[:60]}")

    # 标记处理中
    base_update_status(record_id, VALUE_处理中)

    try:
        # 根据 URL 类型分发
        text = ""
        source = "网页提取"
        url_lower = url.lower()

        if "xiaohongshu.com" in url_lower:
            source = "小红书"
            print(f"📕 小红书内容: {url}")
            page_text, video_url = fetch_xiaohongshu_content(url)
            text = page_text
            if video_url:
                print("  🎬 转写视频...")
                try:
                    transcription = transcribe_video(video_url)
                    if transcription:
                        separator = "\n\n=== 视频转写内容 ===\n\n"
                        text = (text + separator + transcription) if text else transcription
                except Exception as e:
                    print(f"  ⚠️ 视频转写失败: {e}")

        elif "feishu.cn" in url_lower or "larksuite.com" in url_lower:
            source = "飞书文档"
            text = fetch_feishu_doc(url)

        elif any(v in url_lower for v in [".mp4", "bilibili.com", "youtube.com", "douyin.com", "v.qq.com"]):
            source = "视频转写"
            text = transcribe_video(url)

        else:
            text = fetch_web_content(url)

        if not text or not text.strip():
            raise ValueError("未能提取到任何内容")

        print(f"  📊 共 {len(text)} 字符，AI 分析...")
        result = analyze_hr_content(text)

        # 准备更新字段
        update_fields = {
            "标题": result["title"],
            "来源": result.get("source_type", source),
            "整理内容": result["organized_content"][:60000],
            "标签": result["tags"][:3],
            "原始内容": text[:60000],
            FIELD_状态: VALUE_已完成,
        }

        if base_update_record(record_id, update_fields):
            print(f"  ✅ {result['title']} → 已完成")
            return True
        else:
            base_update_status(record_id, VALUE_失败)
            return False

    except Exception as e:
        print(f"  ❌ {type(e).__name__}: {str(e)[:300]}")
        base_update_status(record_id, VALUE_失败)
        return False


# ============================================================
# Daemon / Sync 模式
# ============================================================
_shutdown_flag = False


def _handle_shutdown(signum, frame):
    global _shutdown_flag
    _shutdown_flag = True
    print("\n⏳ 正在停止...")


def run_sync() -> bool:
    """单次处理所有待处理记录后退出。适合 Windows 计划任务"""
    print("🔁 Sync 模式: 处理所有待处理记录...")

    # 孤儿检测：重设卡在"处理中"的记录
    try:
        resp = _run_lark_cli([
            "base", "+record-list",
            "--base-token", BASE_TOKEN,
            "--table-id", TABLE_ID,
            "--limit", "50",
            "--format", "json",
        ], timeout=20)
        records = _parse_record_list(resp)
        for r in records:
            if r.get("fields", {}).get(FIELD_状态) == VALUE_处理中:
                base_update_status(r["id"], VALUE_待处理, "从上一次运行恢复")
                print(f"  🔄 恢复孤儿记录: {r['id'][:8]}...")
    except Exception as e:
        print(f"  ⚠️ 孤儿检测失败: {e}")

    pending = base_get_pending_records(max_records=50)
    if not pending:
        print("✅ 没有待处理的记录")
        return True

    success = fail = 0
    for i, record in enumerate(pending, 1):
        print(f"[{i}/{len(pending)}]", end=" ")
        if process_base_record(record):
            success += 1
        else:
            fail += 1

    print(f"\n{'='*40}")
    print(f"Sync 完成: {success} 成功, {fail} 失败, 共 {len(pending)} 条")
    return fail == 0


def run_daemon(poll_interval: int = 60):
    """常驻后台轮询 Base 中的待处理记录"""
    signal.signal(signal.SIGINT, _handle_shutdown)
    signal.signal(signal.SIGTERM, _handle_shutdown)

    print(f"🔄 Daemon 启动，每 {poll_interval}s 轮询一次")
    print("  在 Base 中创建记录 → 填入链接 → 保存即可自动处理")
    print("  按 Ctrl+C 停止\n")

    while not _shutdown_flag:
        try:
            records = base_get_pending_records()
            if records:
                print(f"📥 发现 {len(records)} 条待处理记录")
                for i, record in enumerate(records, 1):
                    if _shutdown_flag:
                        break
                    print(f"[{i}/{len(records)}]", end=" ")
                    process_base_record(record)
            else:
                # 心跳
                now = datetime.now().strftime("%H:%M:%S")
                print(f"[{now}] .", flush=True)

        except Exception as e:
            print(f"\n⚠️ Daemon 异常: {e}")

        # 逐秒等待，支持快速退出
        for _ in range(poll_interval):
            if _shutdown_flag:
                break
            time.sleep(1)

    print("🛑 Daemon 已停止")

# ============================================================
# 入库流程（含预览/查重）
# ============================================================
def process_and_write(title_hint: str, text: str, source: str, url_link: str,
                      review: bool, skip_confirm: bool) -> bool:
    """通用处理流程：AI分析 → 可选预览 → 查重 → 入库"""
    print("  🤖 AI 分析...")
    result = analyze_hr_content(text)
    final_title = title_hint or result["title"]
    final_source = result.get("source_type", source)

    print(f"  📌 标题: {final_title}")
    print(f"  📂 来源: {final_source}")
    print(f"  🏷️  标签: {', '.join(result['tags'])}")
    print(f"  📝 整理内容: {result['organized_content'][:120]}...")

    # 交互式预览
    if review and not skip_confirm:
        edited = interactive_review(result)
        if edited is None:
            print("  ⏭️ 已跳过")
            return False
        result = edited
        final_title = title_hint or result["title"]
        final_source = result.get("source_type", source)

    # 查重
    if not skip_confirm:
        is_dup = check_duplicate(final_title, text, url_link)
        if is_dup is True:
            print(f"\n  ⚠️ 发现相似内容，是否仍要入库？")
            choice = input("  [y] 仍入库  [s] 跳过: ").strip().lower()
            if choice != "y":
                print("  ⏭️ 已跳过")
                return False

    # 入库
    return write_to_hr_base(
        title=final_title, source=final_source, url_link=url_link,
        organized_content=result["organized_content"], tags=result["tags"],
        raw_text=text
    )

# ============================================================
# 各入口处理函数
# ============================================================
def ingest_file(file_path: str, review: bool, skip_confirm: bool) -> bool:
    print(f"\n📄 处理: {file_path}")
    if not os.path.exists(file_path):
        print(f"  ❌ 文件不存在"); return False

    print("  📖 读取...")
    try:
        text = extract_text(file_path)
    except Exception as e:
        print(f"  ❌ 读取失败: {e}"); return False

    if not text.strip():
        print("  ⚠️ 空内容，跳过"); return False
    print(f"  📊 {len(text)} 字符")

    return process_and_write(
        title_hint="", text=text, source="手动录入", url_link="",
        review=review, skip_confirm=skip_confirm
    )

def ingest_folder(folder_path: str, review: bool, skip_confirm: bool) -> dict:
    supported = (".txt", ".md", ".markdown", ".docx", ".pdf")
    files = []
    for root, _, filenames in os.walk(os.path.abspath(folder_path)):
        for f in filenames:
            if f.lower().endswith(supported):
                files.append(os.path.join(root, f))
    if not files:
        print("📁 无支持的文件"); return {"total": 0, "success": 0, "fail": 0}

    print(f"\n📁 {folder_path} — {len(files)} 个文件\n")
    success = fail = 0
    for i, f in enumerate(files, 1):
        print(f"[{i}/{len(files)}]", end=" ")
        if ingest_file(f, review=review, skip_confirm=skip_confirm):
            success += 1
        else:
            fail += 1
    return {"total": len(files), "success": success, "fail": fail}

def ingest_text(text: str, title: str = "", review: bool = False, skip_confirm: bool = False) -> bool:
    print(f"\n📝 处理文本 ({len(text)} 字符)")
    return process_and_write(
        title_hint=title, text=text, source="手动录入", url_link="",
        review=review, skip_confirm=skip_confirm
    )

def ingest_url(url: str, review: bool, skip_confirm: bool) -> bool:
    # 小红书链接走专用流水线（浏览器渲染+视频转写）
    if "xiaohongshu.com" in url.lower():
        return ingest_xiaohongshu(url, review, skip_confirm)

    src = "网页提取"
    if "mp.weixin.qq.com" in url.lower():
        src = "网页提取"
    print(f"\n🌐 {url}")

    try:
        text = fetch_web_content(url)
    except ValueError as e:
        print(f"  ❌ {e}")
        return False

    if not text.strip():
        print("  ❌ 未能提取到正文"); return False

    return process_and_write(
        title_hint="", text=text, source=src, url_link=url,
        review=review, skip_confirm=skip_confirm
    )

def ingest_video(url: str, review: bool, skip_confirm: bool) -> bool:
    print(f"\n🎬 视频: {url}")
    src = "小红书" if "xiaohongshu.com" in url.lower() else "视频转写"

    try:
        text = transcribe_video(url)
    except (RuntimeError, SystemExit) as e:
        print(f"  ❌ {e}")
        return False

    if not text.strip():
        print("  ❌ 转写结果为空"); return False
    print(f"  📊 {len(text)} 字符")

    return process_and_write(
        title_hint="", text=text, source=src, url_link=url,
        review=review, skip_confirm=skip_confirm
    )

def ingest_feishu_doc(url: str, review: bool, skip_confirm: bool) -> bool:
    print(f"\n📄 飞书文档: {url}")

    try:
        text = fetch_feishu_doc(url)
    except ValueError as e:
        print(f"  ❌ {e}")
        return False

    if not text.strip():
        print("  ❌ 文档内容为空"); return False
    print(f"  📊 {len(text)} 字符")

    return process_and_write(
        title_hint="", text=text, source="飞书文档", url_link=url,
        review=review, skip_confirm=skip_confirm
    )

# ============================================================
# CLI
# ============================================================
def main():
    parser = argparse.ArgumentParser(description="HR知识库 · 内容入库工具（全功能版）")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--file", "-f", help="单个文件")
    group.add_argument("--folder", "-d", help="文件夹(批量)")
    group.add_argument("--url", "-u", help="网页/文章链接")
    group.add_argument("--video", help="视频链接（自动下载+转文字）")
    group.add_argument("--feishu-doc", help="飞书文档/知识库链接")
    group.add_argument("--text", "-t", help="直接输入文字")
    group.add_argument("--daemon", action="store_true", help="Daemon模式：轮询Base自动处理")
    group.add_argument("--sync", action="store_true", help="Sync模式：处理所有待处理记录后退出")
    parser.add_argument("--title", help="手动指定标题（配合 --text/--url 使用）")
    parser.add_argument("--review", "-r", action="store_true", help="强制交互式预览模式")
    parser.add_argument("--yes", action="store_true", help="跳过所有确认（批量/定时任务用）")
    parser.add_argument("--poll-interval", type=int, default=DAEMON_POLL_INTERVAL,
                        help=f"Daemon轮询间隔秒数 (默认 {DAEMON_POLL_INTERVAL})")
    args = parser.parse_args()

    # Daemon / Sync 模式（不需要其他参数）
    if args.daemon:
        run_daemon(poll_interval=args.poll_interval)
        sys.exit(0)
    if args.sync:
        sys.exit(0 if run_sync() else 1)

    # review 默认值：单文件/文本/链接时启用，文件夹批量时关闭；--yes 时关闭
    if args.yes:
        review = False
        skip_confirm = True
    else:
        skip_confirm = False
        review = args.review if args.review else (args.folder is None)

    if args.file:
        sys.exit(0 if ingest_file(args.file, review=review, skip_confirm=skip_confirm) else 1)
    elif args.folder:
        r = ingest_folder(args.folder, review=review, skip_confirm=skip_confirm)
        print(f"\n{'='*40}\n完成: {r['success']}/{r['total']} 成功")
        sys.exit(0 if r['fail'] == 0 else 1)
    elif args.url:
        sys.exit(0 if ingest_url(args.url, review=review, skip_confirm=skip_confirm) else 1)
    elif args.video:
        sys.exit(0 if ingest_video(args.video, review=review, skip_confirm=skip_confirm) else 1)
    elif args.feishu_doc:
        sys.exit(0 if ingest_feishu_doc(args.feishu_doc, review=review, skip_confirm=skip_confirm) else 1)
    elif args.text:
        sys.exit(0 if ingest_text(args.text, args.title or "", review=review, skip_confirm=skip_confirm) else 1)

if __name__ == "__main__":
    main()
