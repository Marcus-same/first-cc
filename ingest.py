"""
AI小组知识库 - 内容入库工具
用法:
  python ingest.py --file "d:/资料/xxx.pdf"
  python ingest.py --file "d:/资料/xxx.docx"
  python ingest.py --folder "d:/资料/AI合集/"
  python ingest.py --url "https://..."
  python ingest.py --text "手动输入的文字内容" --title "标题"

环境变量:
  DEEPSEEK_API_KEY  DeepSeek API Key (必需)
  DEEPSEEK_BASE_URL DeepSeek API 地址 (默认 https://api.deepseek.com)
"""
import argparse, base64, json, mimetypes, os, subprocess, sys, textwrap, time, glob as g
from datetime import datetime
from pathlib import Path
from urllib.request import Request, urlopen
from urllib.error import URLError

# ============================================================
# 配置
# ============================================================
BASE_TOKEN = "KCXjbOD2bafH9Us4hNIcXIRWnLd"
TABLE_ID = "tblseNb1pQVBMnGW"  # 资料库
DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "")
DEEPSEEK_BASE_URL = os.environ.get("DEEPSEEK_BASE_URL", "https://api.deepseek.com")

TAG_POOL = [
    "AI工具", "Prompt技巧", "Dify", "Coze", "教学场景",
    "行业案例", "效率工具", "设计灵感", "技术架构", "其他"
]

# ============================================================
# 文件解析
# ============================================================
def read_txt(path: str) -> str:
    for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()

def read_md(path: str) -> str:
    return read_txt(path)

def read_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            parts.append(row_text)
    return "\n".join(parts)

def read_pdf(path: str) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(path)
        parts = [page.get_text() for page in doc]
        doc.close()
        return "\n\n".join(parts)
    except ImportError:
        raise SystemExit("需要 PyMuPDF: pip install pymupdf")

def read_image(path: str) -> str:
    raise NotImplementedError("图片识别需多模态 LLM，请先手动转换或使用 --text 参数")

def extract_text(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    parsers = {
        ".txt": read_txt, ".md": read_md, ".markdown": read_md,
        ".docx": read_docx, ".pdf": read_pdf,
        ".png": read_image, ".jpg": read_image, ".jpeg": read_image,
    }
    parser = parsers.get(ext)
    if not parser:
        raise ValueError(f"不支持的文件类型: {ext}")
    return parser(file_path)

# ============================================================
# DeepSeek API
# ============================================================
def call_deepseek(prompt: str, max_tokens: int = 1200) -> str:
    if not DEEPSEEK_API_KEY:
        raise SystemExit("请设置环境变量 DEEPSEEK_API_KEY")

    data = json.dumps({
        "model": "deepseek-chat",
        "messages": [
            {"role": "system", "content": "你是一个知识管理助手。你的输出必须严格使用中文。只返回JSON，不要有其他内容。"},
            {"role": "user", "content": prompt}
        ],
        "max_tokens": max_tokens,
        "temperature": 0.3,
        "response_format": {"type": "json_object"}
    }).encode("utf-8")

    req = Request(f"{DEEPSEEK_BASE_URL}/v1/chat/completions", data=data, headers={
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    })
    try:
        resp = urlopen(req, timeout=120)
        result = json.loads(resp.read().decode("utf-8"))
        return result["choices"][0]["message"]["content"]
    except Exception as e:
        raise SystemExit(f"DeepSeek API 调用失败: {e}")

def summarize_and_tag(content: str) -> dict:
    """调用DeepSeek做摘要和打标签，截断过长内容"""
    MAX_LEN = 8000
    if len(content) > MAX_LEN:
        content = content[:MAX_LEN] + "\n\n[... 内容过长，已截断]"

    prompt = f"""请分析以下内容，返回一个JSON对象（不要有任何其他文字）：
{{
  "title": "内容标题（15字以内，概括核心主题）",
  "summary": "内容摘要（150-200字，抓住核心观点和亮点）",
  "key_points": ["关键点1", "关键点2", "关键点3"],
  "tags": ["标签1", "标签2"]
}}

标签必须从以下列表中选择：{', '.join(TAG_POOL)}

你可以从列表中选择1-3个最匹配的标签。

=== 待分析内容 ===
{content}"""

    try:
        result = call_deepseek(prompt)
        parsed = json.loads(result)
        # 验证标签
        valid_tags = [t for t in parsed.get("tags", []) if t in TAG_POOL]
        if not valid_tags:
            valid_tags = ["其他"]
        return {
            "title": parsed.get("title", "未命名"),
            "summary": parsed.get("summary", ""),
            "key_points": parsed.get("key_points", []),
            "tags": valid_tags[:3]
        }
    except json.JSONDecodeError:
        return {"title": "未命名", "summary": content[:200], "key_points": [], "tags": ["其他"]}

# ============================================================
# lark-cli 入库
# ============================================================
def write_to_base(title: str, source: str, url_link: str, summary: str,
                  tags: list[str], full_text: str, file_path: str = "") -> bool:
    """调用 lark-cli 写入资料库"""
    today = datetime.now().strftime("%Y-%m-%d")

    # 标签转成 JSON 数组
    tag_values = ",".join(f'{{"value_type":"text","value":"{t}"}}' for t in tags)
    field_values = [
        {"field_name": "标题", "value": [{"value_type": "text", "value": title}]},
        {"field_name": "来源", "value": [{"value_type": "text", "value": source}]},
        {"field_name": "内容摘要", "value": [{"value_type": "text", "value": summary}]},
        {"field_name": "入库日期", "value": [{"value_type": "text", "value": today}]},
    ]
    if url_link:
        field_values.append({"field_name": "原始链接", "value": [{"value_type": "text", "value": url_link}]})
    if tags:
        field_values.append({"field_name": "关键标签", "value": [
            {"value_type": "text", "value": t} for t in tags
        ]})
    if full_text:
        # 飞书text字段限制，截断到约60000字符
        truncated = full_text[:60000]
        field_values.append({"field_name": "完整文字", "value": [{"value_type": "text", "value": truncated}]})

    body = json.dumps({
        "table_name": "资料库",
        "field_values": field_values
    }, ensure_ascii=False)

    cmd = [
        "lark-cli", "base", "+record-upsert",
        "--base-token", BASE_TOKEN,
        "--table-id", TABLE_ID,
        "--json", body
    ]

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        resp = json.loads(result.stdout) if result.stdout.strip() else {}
        if resp.get("ok"):
            record_id = resp.get("data", {}).get("record", {}).get("id", "?")
            print(f"  ✅ 已入库: {title} (record: {record_id})")
            return True
        else:
            error_msg = resp.get("error", {}).get("message", result.stderr or "未知错误")
            print(f"  ❌ 入库失败: {error_msg}")
            return False
    except Exception as e:
        print(f"  ❌ 调用 lark-cli 失败: {e}")
        return False

# ============================================================
# 单文件处理
# ============================================================
def ingest_file(file_path: str) -> bool:
    print(f"\n📄 处理文件: {file_path}")
    if not os.path.exists(file_path):
        print(f"  ❌ 文件不存在")
        return False

    fname = os.path.basename(file_path)
    ext = Path(file_path).suffix.lower()

    # 读取内容
    print("  📖 读取内容...")
    try:
        text = extract_text(file_path)
    except Exception as e:
        print(f"  ❌ 读取失败: {e}")
        return False

    if not text.strip():
        print("  ⚠️ 文件内容为空，跳过")
        return False
    print(f"  📊 提取 {len(text)} 字符")

    # AI 处理
    print("  🤖 AI 摘要和打标签...")
    result = summarize_and_tag(text)

    # 显示结果
    print(f"  📌 标题: {result['title']}")
    print(f"  🏷️  标签: {', '.join(result['tags'])}")
    print(f"  📝 摘要: {result['summary'][:100]}...")
    if result["key_points"]:
        for pt in result["key_points"]:
            print(f"     • {pt}")

    # 入库
    return write_to_base(
        title=result["title"],
        source="本地文件",
        url_link="",
        summary=result["summary"],
        tags=result["tags"],
        full_text=text,
        file_path=file_path
    )

def ingest_folder(folder_path: str, recursive: bool = True) -> dict:
    """批量导入文件夹"""
    supported = (".txt", ".md", ".markdown", ".docx", ".pdf")
    files = []
    folder_path = os.path.abspath(folder_path)

    if recursive:
        for root, _, filenames in os.walk(folder_path):
            for f in filenames:
                if f.lower().endswith(supported):
                    files.append(os.path.join(root, f))
    else:
        for f in os.listdir(folder_path):
            full = os.path.join(folder_path, f)
            if os.path.isfile(full) and f.lower().endswith(supported):
                files.append(full)

    if not files:
        print(f"📁 {folder_path} 中没有找到支持的文档文件")
        return {"total": 0, "success": 0, "fail": 0}

    print(f"\n📁 文件夹: {folder_path}")
    print(f"   找到 {len(files)} 个文档文件\n")

    success, fail = 0, 0
    for i, f in enumerate(files, 1):
        print(f"[{i}/{len(files)}]", end=" ")
        if ingest_file(f):
            success += 1
        else:
            fail += 1
        if i < len(files):
            print()  # 分隔空行

    return {"total": len(files), "success": success, "fail": fail}

# ============================================================
# URL 处理 (简易版，后续扩展)
# ============================================================
def ingest_url(url: str) -> bool:
    print(f"\n🌐 处理链接: {url}")
    print("  ⚠️ URL内容抓取功能开发中，请先用 --text 参数手动输入内容")
    return False

# ============================================================
# 纯文本处理
# ============================================================
def ingest_text(text: str, title: str = "") -> bool:
    print(f"\n📝 处理文本 ({len(text)} 字符)")
    print("  🤖 AI 摘要和打标签...")
    result = summarize_and_tag(text)
    final_title = title or result["title"]

    print(f"  📌 标题: {final_title}")
    print(f"  🏷️  标签: {', '.join(result['tags'])}")
    return write_to_base(
        title=final_title, source="其他", url_link="",
        summary=result["summary"], tags=result["tags"],
        full_text=text
    )

# ============================================================
# CLI
# ============================================================
def main():
    parser = argparse.ArgumentParser(description="AI小组知识库 · 内容入库工具")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--file", "-f", help="单个文件路径")
    group.add_argument("--folder", "-d", help="文件夹路径（批量导入）")
    group.add_argument("--url", "-u", help="网页/小红书链接")
    group.add_argument("--text", "-t", help="直接输入文字内容")
    parser.add_argument("--title", help="手动指定标题（配合 --text 使用）")
    parser.add_argument("--no-recursive", action="store_true", help="不递归遍历子文件夹")
    args = parser.parse_args()

    if args.file:
        ok = ingest_file(args.file)
        sys.exit(0 if ok else 1)
    elif args.folder:
        result = ingest_folder(args.folder, recursive=not args.no_recursive)
        print(f"\n{'='*40}")
        print(f"完成: {result['success']} 成功, {result['fail']} 失败, 共 {result['total']} 个文件")
        sys.exit(0 if result['fail'] == 0 else 1)
    elif args.url:
        ok = ingest_url(args.url)
        sys.exit(0 if ok else 1)
    elif args.text:
        ok = ingest_text(args.text, args.title or "")
        sys.exit(0 if ok else 1)

if __name__ == "__main__":
    main()
